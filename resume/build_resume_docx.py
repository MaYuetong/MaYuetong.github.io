"""
Builds resume.docx matching Yuetong's original CV PDF style.
Run: python3 build_resume_docx.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "resume.docx")
BLACK = RGBColor(0x1d,0x1d,0x1f)
GREY  = RGBColor(0x55,0x55,0x55)

def sf(run, size, bold=False, italic=False, colour=BLACK):
    run.font.name="Calibri"; run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic; run.font.color.rgb=colour

def p(doc, sb=0, sa=3, indent=0):
    pg=doc.add_paragraph()
    pg.paragraph_format.space_before=Pt(sb)
    pg.paragraph_format.space_after=Pt(sa)
    pg.paragraph_format.line_spacing=Pt(14)
    if indent: pg.paragraph_format.left_indent=Cm(indent)
    return pg

def section(doc, title):
    pg=p(doc, sb=8, sa=2)
    r=pg.add_run(title); sf(r,11,bold=True)
    # bottom border
    pPr=pg._p.get_or_add_pPr()
    pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"000000")
    pBdr.append(bot); pPr.append(pBdr)

def exp_role(doc, company, role, date, location=None):
    # Company bold
    pg=p(doc, sb=5, sa=1)
    sf(pg.add_run(company), 10.5, bold=True)
    # Role italic + date
    pg2=p(doc, sb=0, sa=1)
    sf(pg2.add_run(role+" - "), 10, italic=True)
    if location:
        sf(pg2.add_run(location+" | "), 10, italic=True, colour=GREY)
    sf(pg2.add_run(date), 10, italic=True, colour=GREY)

def bullets(doc, items):
    for i,b in enumerate(items,1):
        pg=p(doc, sb=0, sa=2, indent=0.4)
        sf(pg.add_run(f"{i}) {b}"), 10)

def edu_entry(doc, school, location, dates, degree, modules):
    pg=p(doc, sb=4, sa=1)
    r=pg.add_run("• "); sf(r,10.5,bold=True)
    sf(pg.add_run(school), 10.5, bold=True)
    sf(pg.add_run(f" – {location} ({dates})"), 10)
    pg2=p(doc, sb=0, sa=1, indent=0.4)
    sf(pg2.add_run(degree), 10, italic=True)
    pg3=p(doc, sb=0, sa=2, indent=0.4)
    sf(pg3.add_run("Modules: "+modules), 10)

doc=Document()
for s in doc.sections:
    s.top_margin=s.bottom_margin=Cm(2.0)
    s.left_margin=s.right_margin=Cm(2.5)
doc.styles["Normal"].font.name="Calibri"
doc.styles["Normal"].font.size=Pt(10.5)

# ── NAME ──
pg=p(doc,sb=0,sa=2)
sf(pg.add_run("Yuetong Ma"),14,bold=True)

# ── CONTACT ──
for line in [
    "Email: joannasoo@163.com",
    "Spatial CV (Portfolio): https://arcg.is/j9qHv",
    "LinkedIn: https://www.linkedin.com/in/yuetong-ma-a2258324b/",
    "Personal Website: [add hosted URL]",
    "Languages: Chinese (Native), English (Professional), German (Basic)",
]:
    pg=p(doc,sb=0,sa=1); sf(pg.add_run(line),10)

# ── EDUCATION ──
section(doc,"Education")
edu_entry(doc,"Peking University","Beijing, China","09/2021 – 07/2024",
    "MSc in Geography - GPA: 3.74/4.0 (with distinction)",
    "Geosoft Computing Models and Applications (WebGIS; Python); Quantitative Methods in Geography; Machine Learning (Python; Matlab); Big Data Analysis (Python; ArcGIS); Land Use and Policy.")
edu_entry(doc,"Stockholm University","Stockholm, Sweden","09/2022 – 03/2023",
    "Exchange Program",
    "Human Geography; Communication; Globalisation and Strategy.")
edu_entry(doc,"Beijing Forestry University","Beijing, China","09/2017 – 07/2021",
    "BSc in Forestry (Urban Forestry) - GPA: 91.22/100 (with distinction)",
    "Mathematical Statistics (Python; R); Urban Ecology; Forest Ecology; Landscape Ecology; Forestry Remote Sensing and Geographic Information System (GIS).")

# ── WORK EXPERIENCE ──
section(doc,"Work Experience")

exp_role(doc,
    "United Nations Secretariat – Department of Peace Operations (DPO), UNOCC",
    "Geopolitical Analyst (UN Trainee)","11/2025 – 05/2026","Hybrid (New York-based)")
bullets(doc,[
    "Designed an AI Agent-based system to automatically identify geospatial information needs related to peace and political affairs.",
    "Developed an automated \"demand–task–operation\" pipeline converting unstructured requests into executable GIS tasks.",
])

exp_role(doc,
    "DJI Innovation (Shenzhen DJI Sciences and Technologies Ltd.)",
    "Product Manager – Technical Service Solutions","02/2025 – 03/2026","Shenzhen, China")
bullets(doc,[
    "Led end-to-end design of an automated service solution delivery workflow, including communication protocols and service dashboards. Improved service response efficiency by 40% and increased automation rate from 0% to 75%. Used enterprise systems (CRM; LMS; RMS) to support service operations and data workflows.",
    "Contributed to the development of an AI Agent-based service platform integrating AI customer service and automated damage assessment.",
    "Built a multi-objective optimization model for the North America warehouse network and proposed cost-saving scenarios, with an estimated 30% reduction in logistics costs.",
])

exp_role(doc,
    "International Atomic Energy Agency (IAEA)",
    "Regional GIS Trainer and Technical Expert (SSA)","08/2025","Faisalabad, Pakistan")
bullets(doc,[
    "Designed and delivered training on open-source geo-visualization tools for monitoring vector-borne diseases and animal health. Prepared technical documentation and end-of-mission report in accordance with IAEA standards.",
    "Developed training materials covering data preparation, CRS, symbology, heatmaps, buffers, grids, animations and map publishing. Led hands-on sessions for government officials and researchers from the Middle East, Africa and Southeast Asia.",
])

exp_role(doc,
    "International Atomic Energy Agency (IAEA)",
    "Trainee – WebGIS Developer & Data Analyst / Project Assistant","01/2024 – 01/2025","Vienna, Austria")
bullets(doc,[
    "Developed a GIS-based animal health monitoring platform enabling data upload, management and geospatial visualization. Integrated spatial and temporal datasets to improve efficiency of zoonotic disease surveillance workflows.",
    "Produced thematic maps and visual products for IAEA General Conference and annual reports. Drafted project progress and results reports highlighting best practices and impact.",
    "Developed and delivered GIS training materials for agricultural and veterinary applications.",
])

exp_role(doc,
    "UNESCO – WHITRAP",
    "Research Assistant – Climate Change and Heritage Conservation","12/2022 – 04/2023","Remote (Sweden-based)")
bullets(doc,[
    "Contributed to the Metropolitan Heritage Project by assessing climate change impacts on heritage conservation.",
    "Conducted interdisciplinary research combining heritage studies, digitalization and geospatial methods.",
    "Research on AI and geographic approaches to analyse climate risks and resilience. Prepared literature reviews and research reports for project stakeholders.",
])

exp_role(doc,
    "China International Engineering Consulting Corporation (CIECC)",
    "Project Assistant (Internship) – Infrastructure Investment Department","01/2021 – 06/2021","Beijing, China")
bullets(doc,[
    "Contributed to the feasibility study of the China National Botanical Garden through systematic policy and literature review. Compiled benchmark indicators of representative botanical gardens for international comparison.",
    "Supported preparation of the 14th Five-Year Plan for Infrastructure Development in Tongzhou District.",
    "Drafted technical documents and presentation materials for government stakeholders.",
])

# ── RESEARCH PROJECTS ──
section(doc,"Research Projects")

for proj, role_date, pts in [
    ("Regional Resilience Assessment Using Night-Time Light Data","Project Lead | 05/2022 – 01/2024",[
        "Analysed NTL data (NASA Black Marble) to assess spatio-temporal resilience patterns during Typhoon Mangkhut. Developed a patented method for dynamic pattern recognition of disaster-affected areas. Applied machine learning (LightGBM) and interpretable AI methods to identify non-linear relationships between socio-environmental factors and disturbance and recovery rates.",
        "Awarded First Prize at the European Cartographic Conference 2024; two SCI Q1 papers under review.",
    ]),
    ("Green View Index (GVI) Analysis of Beijing Central Axis","Project Lead | 04/2019 – 09/2019",[
        "Analysed spatial variation of GVI using street view imagery and GIS techniques. Identified key environmental and urban form factors influencing GVI distribution.",
        "Publication: Ma Yuetong. Analysis of Green View Index in Beijing Central Axis Road System Based on Street View Image. Journal of Chinese Urban Forestry. 2020, 18(03).",
    ]),
    ("Undergraduate Innovation and Entrepreneurship Program","Project Lead | 06/2019 – 06/2021",[
        "Established phenological evaluation criteria based on field observation of over 300 individuals. Analysed phenological and physiological data using statistical methods.",
        "Awarded RMB 10,000 research funding and Excellent Graduation Thesis.",
    ]),
]:
    pg=p(doc,sb=5,sa=1); sf(pg.add_run(proj),10.5,bold=True)
    pg2=p(doc,sb=0,sa=1); sf(pg2.add_run(role_date),10,italic=True,colour=GREY)
    for pt in pts:
        pg3=p(doc,sb=0,sa=2,indent=0.4); sf(pg3.add_run("• "+pt),10)

# ── VISION ──
section(doc,"Development Direction & Vision")
pg=p(doc,sb=3,sa=4)
sf(pg.add_run(
    "My professional focus is on building decision-support systems based on spatial data and artificial intelligence. "
    "I aim to integrate geospatial intelligence and AI into operational platforms that support policy-making, "
    "crisis response and organizational efficiency."
),10)

doc.save(OUT)
print(f"Saved: {OUT}")
